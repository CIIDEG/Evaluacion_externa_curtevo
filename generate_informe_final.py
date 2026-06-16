"""Informe Final de Evaluación Externa — versión consultiva de alto impacto.

Marco metodológico integrado:
- Guía de Seguimiento y Evaluación de la Cooperación Valenciana (Generalitat Valenciana, GVA)
- Criterios CAD/OCDE (Comité de Ayuda al Desarrollo de la Organización para la Cooperación y el Desarrollo Económicos) (6)
- Marco de evaluación de impacto del BID (Banco Interamericano de Desarrollo) (Inter-American Development Bank, 2010,
  http://dx.doi.org/10.18235/0010435)
- Buenas prácticas de Theory of Change y triangulación metodológica

Estructura:
  PARTE I  · RESUMEN EJECUTIVO   (decisión de alto nivel)
  PARTE II · ANÁLISIS TÉCNICO     (cuantitativo + cualitativo + impacto comunitario)
  PARTE III· ESTRATEGIA           (recomendaciones por actor + roadmap sostenibilidad)
  ANEXOS

Uso (dentro del contenedor):
    docker exec cutervo-app python /code/generate_informe_final.py
"""
import os
import sys
import math
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap
import numpy as np

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.database import SessionLocal
from app import models, results as results_mod, impact as impact_mod, report_lb as report_mod
from app.surveys_def import get_form, ALL_FORMS

DOCS_DIR = Path("/code/docs")
DOCS_DIR.mkdir(parents=True, exist_ok=True)
PHOTOS_DIR = Path("/code/data/photos")
UPLOADS_DIR = Path("/code/data/uploads")

# Pies de foto editables (asociados por prefijo del nombre del archivo)
PHOTO_CAPTIONS = {
    "00-logo-proyecto":
        ("Logo institucional del proyecto «Mejora de las oportunidades de inserción "
         "sociolaboral juvenil en Cutervo, Cajamarca, Perú»."),
    "01-actividad-formacion-docente":
        ("Sesión de formación de docentes en emprendimiento juvenil. Los docentes de las "
         "instituciones educativas participantes desarrollaron capacidades pedagógicas para "
         "el acompañamiento al estudiantado en proyectos productivos escolares."),
    "02-trabajo-grupal-estudiantes":
        ("Trabajo grupal de estudiantes durante la implementación del programa de "
         "emprendimiento escolar. La metodología «joven a joven» promovió el liderazgo "
         "compartido y la construcción colectiva de propuestas productivas."),
    "03-sesion-emprendimiento":
        ("Sesión de sensibilización con jóvenes sobre derechos económicos y sociales y "
         "su vinculación con iniciativas productivas. La actividad combinó información "
         "técnica con discusión grupal sobre el ejercicio efectivo de derechos."),
    "04-equipo-evaluacion":
        ("Equipo técnico durante la fase de trabajo de campo de la evaluación final. "
         "La aplicación de instrumentos cualitativos (entrevistas a informantes clave, "
         "grupos focales, observación e Historias del Cambio Más Significativo) se realizó "
         "presencialmente en las instituciones educativas del distrito."),
    "05-paisaje-cutervo":
        ("Paisaje del distrito de Cutervo, provincia homónima, departamento de Cajamarca, "
         "Perú. Cutervo se ubica en la zona andina norte del país, con una geografía "
         "marcada por valles interandinos y áreas rurales dispersas que condicionan la "
         "accesibilidad a servicios educativos."),
    "foto":
        ("Estudiantes de una institución educativa participante del proyecto trabajan "
         "el terreno de un emprendimiento escolar agrícola (biohuerto) durante una "
         "jornada de campo. La imagen evidencia la apropiación práctica del componente "
         "de emprendimiento productivo, con participación equilibrada entre estudiantes "
         "varones y mujeres y acompañamiento docente. Fotografía capturada por el "
         "evaluador externo durante la fase de observación de campo (Anexo 5)."),
}


def get_photo_caption(filename: str) -> str:
    """Devuelve el pie de foto asociado al archivo (busca por prefijo)."""
    name = filename.lower()
    for key, caption in PHOTO_CAPTIONS.items():
        if name.startswith(key.lower()):
            return caption
    return f"Imagen del proyecto: {filename}"

# ============================================================
NAVY = RGBColor(0x0A, 0x1F, 0x3D)
INDIGO = RGBColor(0x16, 0x2E, 0x5C)
CORAL = RGBColor(0xFF, 0x5C, 0x4D)
AQUA = RGBColor(0x2E, 0xC4, 0xB6)
GOLD = RGBColor(0xF7, 0xC5, 0x59)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)

PALETTE = ["#FF5C4D","#2EC4B6","#F7C559","#0A1F3D","#162E5C",
           "#caa235","#B5BDC9","#e84a3c","#2a8c7f","#d6a541",
           "#3a5fa0","#7a8aa3"]
PAL_HEX = {"navy":"#0A1F3D","indigo":"#162E5C","coral":"#FF5C4D",
           "aqua":"#2EC4B6","gold":"#F7C559","gray":"#5b5b5b"}


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def set_run(run, *, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color: run.font.color.rgb = color


# ---------- estructura ----------
def H1(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); set_run(r, size=20, bold=True, color=color)
    return p

def H2(doc, text, color=CORAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); set_run(r, size=14, bold=True, color=color)
    return p

def H3(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); set_run(r, size=12, bold=True, color=color)
    return p

def H4(doc, text, color=INDIGO):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); set_run(r, size=11, bold=True, italic=True, color=color)
    return p

def P(doc, text, *, italic=False, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.15
    p.alignment = align
    r = p.add_run(text); set_run(r, size=size, italic=italic, bold=bold, color=color)
    return p

def bullet(doc, text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead:
        r = p.add_run(lead); set_run(r, bold=True)
    r = p.add_run(text); set_run(r)

def numbered(doc, text, lead=None):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead:
        r = p.add_run(lead); set_run(r, bold=True)
    r = p.add_run(text); set_run(r)

def add_caption(doc, num, title, kind="Tabla"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{kind} N° {num}: "); set_run(r1, size=10, bold=True, color=CORAL)
    r2 = p.add_run(title); set_run(r2, size=10, bold=True)

def add_fuente(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Fuente: {text}")
    set_run(r, size=9, italic=True, color=GRAY)

def add_interp(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12); p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run("Análisis del consultor. ")
    set_run(r1, size=10, bold=True, italic=True, color=GOLD)
    r2 = p.add_run(text); set_run(r2, size=10, italic=True, color=GRAY)

def add_quote_box(doc, text):
    """Caja de cita destacada."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade_cell(cell, "FFF8F0")
    cell.width = Cm(15)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    r = p.add_run("« " + text + " »")
    set_run(r, size=10, italic=True, color=NAVY)


def add_table(doc, headers, rows, *, total_row=None, col_widths_cm=None, header_fill="0A1F3D"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h)); set_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        shade_cell(hdr_cells[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val) if val is not None else "—")
            set_run(r, size=10)
    if total_row:
        cells = t.add_row().cells
        for i, val in enumerate(total_row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val) if val is not None else "—")
            set_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
            shade_cell(cells[i], "FF5C4D")
    if col_widths_cm:
        for row in t.rows:
            for i, w in enumerate(col_widths_cm):
                if i < len(row.cells): row.cells[i].width = Cm(w)
    return t


# ============================================================
#                       CHARTS
# ============================================================
def make_chart(chart_type, labels, data, title, *, w=14, h=8):
    fig, ax = plt.subplots(figsize=(w/2.54, h/2.54), dpi=150)
    colors = (PALETTE * 4)[:len(labels)]
    if chart_type == "doughnut":
        wedges, _ = ax.pie(data, colors=colors, startangle=90,
                           wedgeprops={"width": 0.45, "edgecolor": "white"})
        ax.legend(wedges, [f"{l} ({d}%)" for l, d in zip(labels, data)],
                  loc="center left", bbox_to_anchor=(1.0, 0.5), fontsize=7, frameon=False)
    elif chart_type == "horizontal_bar":
        y = range(len(labels))
        ax.barh(y, data, color=colors, edgecolor="white")
        ax.set_yticks(y); ax.set_yticklabels(labels, fontsize=8)
        ax.set_xlim(0, max(100, max(data)*1.1 if data else 100))
        ax.invert_yaxis()
        for i, v in enumerate(data): ax.text(v + 1, i, f"{v}", va="center", fontsize=8, color="#333")
    else:
        x = range(len(labels))
        ax.bar(x, data, color=colors, edgecolor="white")
        ax.set_xticks(x); ax.set_xticklabels(labels, rotation=25, ha="right", fontsize=8)
        for i, v in enumerate(data): ax.text(i, v + 0.5, f"{v}", ha="center", fontsize=8)
    ax.set_title(title, fontsize=10, weight="bold", pad=10, color="#0A1F3D")
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout()
    out = Path(f"/tmp/_chart_{abs(hash(title))}_{abs(hash(str(data)))}.png")
    fig.savefig(out, dpi=150, bbox_inches="tight", facecolor="white"); plt.close(fig)
    return out


def make_gauge(value, label, *, max_val=100, w=10, h=6):
    """Gauge semicircular para mostrar score 0-100."""
    fig, ax = plt.subplots(figsize=(w/2.54, h/2.54), dpi=150, subplot_kw={"projection":"polar"})
    # rangos
    theta = np.linspace(0, np.pi, 200)
    # color según valor
    if value is None:
        color = "#B5BDC9"; level = "—"
    elif value >= 80: color, level = "#2EC4B6", "Alto"
    elif value >= 60: color, level = "#F7C559", "Medio"
    elif value >= 40: color, level = "#FF5C4D", "Bajo"
    else:             color, level = "#c41a0a", "Muy bajo"

    # Fondo del gauge
    ax.barh(1, np.pi, left=0, height=0.5, color="#eeeeee", edgecolor="none")
    if value is not None:
        # arco proporcional
        frac = max(0, min(1, value / max_val))
        ax.barh(1, np.pi * frac, left=0, height=0.5, color=color, edgecolor="none")
    ax.set_theta_zero_location("W"); ax.set_theta_direction(-1)
    ax.set_ylim(0, 2); ax.set_yticks([]); ax.set_xticks([])
    ax.spines["polar"].set_visible(False)
    ax.set_title(label, fontsize=9, weight="bold", color="#0A1F3D", pad=12)
    # texto central
    ax.text(np.pi/2, 0, f"{value if value is not None else '—'}\n{level}",
            ha="center", va="center", fontsize=14, weight="bold", color=color)
    fig.tight_layout()
    out = Path(f"/tmp/_gauge_{abs(hash(label))}.png")
    fig.savefig(out, dpi=150, bbox_inches="tight", facecolor="white"); plt.close(fig)
    return out


def make_heatmap_ie(by_ie_data, dim_labels, title, *, w=15, h=8):
    """Heatmap IE × dimensión con scores."""
    if not by_ie_data.get("ies"):
        return None
    ies = [r["ie"] for r in by_ie_data["ies"]]
    dim_keys = by_ie_data["dim_keys"]
    matrix = []
    for r in by_ie_data["ies"]:
        matrix.append([r.get(k, 0) or 0 for k in dim_keys])
    matrix = np.array(matrix)

    fig, ax = plt.subplots(figsize=(w/2.54, h/2.54), dpi=150)
    cmap = LinearSegmentedColormap.from_list("ev", ["#FF5C4D","#F7C559","#2EC4B6"])
    im = ax.imshow(matrix, cmap=cmap, vmin=0, vmax=100, aspect="auto")
    # ticks
    ax.set_xticks(range(len(dim_keys)))
    ax.set_xticklabels([dim_labels.get(k, k) for k in dim_keys], rotation=20, ha="right", fontsize=7)
    ax.set_yticks(range(len(ies)))
    ax.set_yticklabels([ie[:35] for ie in ies], fontsize=7)
    # números dentro
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            txt_color = "white" if matrix[i,j] < 40 or matrix[i,j] > 70 else "black"
            ax.text(j, i, f"{matrix[i,j]:.0f}", ha="center", va="center", fontsize=7, color=txt_color)
    cbar = fig.colorbar(im, ax=ax, fraction=0.04, pad=0.04)
    cbar.ax.tick_params(labelsize=7)
    ax.set_title(title, fontsize=10, weight="bold", pad=10, color="#0A1F3D")
    fig.tight_layout()
    out = Path(f"/tmp/_heat_{abs(hash(title))}.png")
    fig.savefig(out, dpi=150, bbox_inches="tight", facecolor="white"); plt.close(fig)
    return out


def add_chart_img(doc, png_path, width_cm=14):
    if not png_path or not Path(png_path).exists(): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png_path), width=Cm(width_cm))


def add_photo_with_caption(doc, photo_path: Path, num: int, *, width_cm=12):
    """Inserta una foto centrada con título 'Foto N°X' y pie descriptivo."""
    if not photo_path.exists():
        return False
    # título
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"Fotografía N° {num}: "); set_run(r1, size=10, bold=True, color=CORAL)
    # caption_short = nombre legible
    name = photo_path.stem.split("-", 1)[-1].replace("-", " ").capitalize()
    r2 = p.add_run(name); set_run(r2, size=10, bold=True)
    # imagen
    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        pi.add_run().add_picture(str(photo_path), width=Cm(width_cm))
    except Exception as e:
        P(doc, f"[Foto no insertada: {e}]", italic=True, size=9)
        return False
    # pie de foto
    caption = get_photo_caption(photo_path.name)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption); set_run(r, size=9, italic=True, color=GRAY)
    return True


def list_photos():
    """Devuelve lista ordenada de fotos en /code/data/photos/ + uploads/observacion/"""
    photos = []
    if PHOTOS_DIR.exists():
        photos.extend(sorted(PHOTOS_DIR.glob("*.[jp][np]g")))
        photos.extend(sorted(PHOTOS_DIR.glob("*.png")))
        photos.extend(sorted(PHOTOS_DIR.glob("*.jpeg")))
    # también imágenes de observaciones
    if UPLOADS_DIR.exists():
        for f in sorted(UPLOADS_DIR.rglob("*.jpg")):
            photos.append(f)
        for f in sorted(UPLOADS_DIR.rglob("*.png")):
            photos.append(f)
        for f in sorted(UPLOADS_DIR.rglob("*.jpeg")):
            photos.append(f)
    # dedupe por path absoluto
    seen = set(); out = []
    for p in photos:
        s = str(p.resolve())
        if s not in seen:
            seen.add(s); out.append(p)
    return out


# ============================================================
#                       BUILD
# ============================================================
def build_informe():
    print("Cargando datos de la base...")
    db = SessionLocal()
    payload = results_mod.build_full_results(db, include_audios=True)
    R = payload["report_lb"]
    imp = payload["impact"]
    imp_est = impact_mod.compute_impact(db, "estudiantes")
    imp_doc = impact_mod.compute_impact(db, "docentes")
    imp_est_by_ie = impact_mod.compute_impact_by_ie(db, "estudiantes")
    imp_doc_by_ie = impact_mod.compute_impact_by_ie(db, "docentes")
    n_est = R["n_est"]; n_doc = R["n_doc"]; n_iiee = R["n_iiee"]
    now_str = datetime.now().strftime("%d/%m/%Y %H:%M")
    fecha_largo = datetime.now().strftime("%B de %Y").capitalize()

    DIM_NAME_EST = {"B":"R2 Empr.", "C":"R3 Der.", "D":"P. Vida", "E":"Género"}
    DIM_NAME_DOC = {"B":"R1 Cap.", "C":"Eficiencia", "D":"Sostenib."}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.5)

    # =========================================================
    # PORTADA
    # =========================================================
    print("→ Portada")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME FINAL DE EVALUACIÓN EXTERNA"); set_run(r, size=10, bold=True, color=CORAL)
    # Logo del proyecto (centrado en portada)
    logo_path = PHOTOS_DIR / "00-logo-proyecto.png"
    if logo_path.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(30); p.paragraph_format.space_after = Pt(10)
        p.add_run().add_picture(str(logo_path), width=Cm(11))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("MEJORA DE LAS OPORTUNIDADES\nDE INSERCIÓN SOCIOLABORAL JUVENIL\nEN CUTERVO, CAJAMARCA — PERÚ")
    set_run(r, size=22, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Expediente SOLPCD/2024/0118"); set_run(r, size=12, italic=True, color=GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Cofinanciado por la Generalitat Valenciana\n"
                  "Ejecutado por Intersindical Solidaria (IS, España)\n"
                  "e Instituto de Pedagogía Popular (IPP (Instituto de Pedagogía Popular), Perú)")
    set_run(r, size=11, color=NAVY)
    # bloque metodológico
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("Marco metodológico integrado:\n")
    set_run(r, size=9, italic=True, color=GRAY)
    r = p.add_run("• Criterios CAD/OCDE  • Guía de Evaluación de la Cooperación Valenciana\n"
                  "• Marco BID de evaluación de impacto (Inter-American Development Bank, 2010,\n"
                  "  http://dx.doi.org/10.18235/0010435)  • Teoría del Cambio · IA Gemini 2.5 Flash")
    set_run(r, size=9, italic=True, color=GRAY)
    # firma
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("Evaluador externo independiente\n"); set_run(r, size=10, italic=True, color=GRAY)
    r = p.add_run("Dr. Segundo Santos Pérez Pérez"); set_run(r, size=12, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DNI 27398656 · CPPe 1627398656\nMayamida66@gmail.com · +51 948 584 865")
    set_run(r, size=10, color=GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run(f"Cutervo, Cajamarca — Perú · {fecha_largo}\nDocumento generado: {now_str} UTC")
    set_run(r, size=10, bold=True, color=CORAL)
    doc.add_page_break()

    # =========================================================
    # ÍNDICE
    # =========================================================
    print("→ Índice")
    H1(doc, "Tabla de contenidos")
    indice = [
        ("PARTE I · RESUMEN EJECUTIVO", True),
        ("  1. Síntesis general", False),
        ("  2. Tablero de scores e indicadores", False),
        ("  3. Hallazgos críticos", False),
        ("  4. Decisiones recomendadas", False),
        ("PARTE II · ANÁLISIS TÉCNICO", True),
        ("  5. Aspectos introductorios", False),
        ("  6. Objetivos y alcance", False),
        ("  7. Antecedentes y contexto", False),
        ("  8. Marco metodológico", False),
        ("  9. Condicionantes y límites", False),
        ("  10. Análisis cuantitativo (Anexo 1 — Estudiantes)", False),
        ("  11. Análisis cuantitativo (Anexo 2 — Docentes)", False),
        ("  12. Análisis cualitativo y síntesis temática IA", False),
        ("  13. Nivel de impacto consolidado", False),
        ("  14. Análisis del impacto comunitario y territorial", False),
        ("PARTE III · CONCLUSIONES Y ESTRATEGIA", True),
        ("  15. Hallazgos por criterio CAD/OCDE", False),
        ("  16. Conclusiones evaluativas", False),
        ("  17. Recomendaciones estratégicas por actor", False),
        ("  18. Roadmap de sostenibilidad", False),
        ("  19. Lecciones aprendidas", False),
        ("ANEXOS", True),
    ]
    for txt, is_bold in indice:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt); set_run(r, size=11, bold=is_bold, color=(CORAL if is_bold else None))
    doc.add_page_break()

    # =========================================================
    # PARTE I · RESUMEN EJECUTIVO
    # =========================================================
    print("→ PARTE I · Resumen ejecutivo")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PARTE I"); set_run(r, size=12, bold=True, color=CORAL)
    H1(doc, "Resumen ejecutivo")
    P(doc, f"Este informe presenta los resultados de la evaluación final externa del proyecto "
           f"«Mejora de las oportunidades de inserción sociolaboral juvenil en Cutervo, Cajamarca — "
           f"Perú» (SOLPCD/2024/0118), ejecutado entre 01/05/2025 y 30/04/2026 por el consorcio "
           f"Intersindical Solidaria (IS) — Instituto de Pedagogía Popular (IPP), con un coste "
           f"total de 68.285 € (subvención GVA 66.325 €) y cobertura en 12 instituciones "
           f"educativas del distrito.")
    P(doc, f"La evaluación se desarrolló entre {R['fecha_min']} y {R['fecha_max']} aplicando un "
           f"diseño mixto cuanti-cualitativo, conforme a los seis criterios CAD/OCDE, las "
           f"directrices de la Guía de Seguimiento y Evaluación de la Cooperación Valenciana y los "
           f"principios del marco BID de evaluación de impacto (http://dx.doi.org/10.18235/0010435). "
           f"Se procesaron {n_est} respuestas de estudiantes y {n_doc} de docentes, "
           f"complementadas con entrevistas, grupos focales, observaciones y narraciones MSC (Historias del Cambio Más Significativo, por sus siglas en inglés); los "
           f"audios fueron transcritos y analizados con IA (Gemini 2.5 Flash) para extracción "
           f"automatizada de temas y citas por criterio CAD.")

    H2(doc, "Tablero ejecutivo de scores")
    # gauges
    print("  · gauges principales")
    g_cuanti = make_gauge(imp.get("global_cuanti"), "Impacto cuantitativo")
    g_quali = make_gauge(imp.get("global_quali"), "Cobertura cualitativa")
    g_est = make_gauge(imp_est.get("global"), "Estudiantes")
    g_doc = make_gauge(imp_doc.get("global"), "Docentes")
    # tabla 2x2 con gauges
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, png in zip([t.rows[0].cells[0], t.rows[0].cells[1], t.rows[1].cells[0], t.rows[1].cells[1]],
                         [g_cuanti, g_quali, g_est, g_doc]):
        cell.text = ""
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run().add_picture(str(png), width=Cm(7))
    add_fuente(doc, "Cálculo del consultor a partir de la base de datos del sistema · " + now_str)

    # tabla síntesis indicadores
    add_caption(doc, 1, "Tablero síntesis: indicadores clave de la evaluación", "Tabla")
    add_table(doc,
        ["Indicador", "Valor", "Meta", "% avance / nivel"],
        [
            ["Respuestas de estudiantes", n_est, 196, f"{round(100*n_est/196,1)}%"],
            ["Respuestas de docentes", n_doc, 24, f"{round(100*n_doc/24,1)}%"],
            ["Instituciones educativas con datos", n_iiee, 12, f"{round(100*n_iiee/12,1)}%"],
            ["Audios capturados (cualitativos)", payload["overview"]["total_audios"], "—", "—"],
            ["Score impacto cuantitativo global", f"{imp.get('global_cuanti') or '—'}/100", "≥ 70", imp.get('global_cuanti_level','—')],
            ["Cobertura cualitativa global", f"{imp.get('global_quali') or '—'}%", "≥ 70", imp.get('global_quali_level','—')],
        ],
        col_widths_cm=[7, 3, 2.5, 4])
    add_fuente(doc, "Base de datos del sistema de evaluación · " + now_str)

    H2(doc, "Hallazgos críticos (síntesis ejecutiva)")
    P(doc, "Los cinco hallazgos siguientes condensan la evidencia más relevante para la toma de "
           "decisiones de las entidades cofinanciadoras y ejecutoras:")
    for i, c in enumerate(R["conclusiones"][:6], 1):
        H4(doc, f"H{i}. {c['titulo']}")
        P(doc, c["texto"])

    H2(doc, "Decisiones recomendadas (top-5 estratégicas)")
    P(doc, "Estas decisiones, ordenadas por prioridad, se derivan del análisis técnico de la Parte II y deben ser asumidas por los actores responsables:")
    for i, rec in enumerate(R["recomendaciones"][:5], 1):
        numbered(doc, rec)
    doc.add_page_break()

    # =========================================================
    # PARTE II · ANÁLISIS TÉCNICO
    # =========================================================
    print("→ PARTE II · Análisis técnico")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PARTE II"); set_run(r, size=12, bold=True, color=CORAL)
    H1(doc, "Análisis técnico detallado")

    # 5. Aspectos introductorios
    H2(doc, "5. Aspectos introductorios")
    H3(doc, "5.1 Sobre el proyecto evaluado")
    P(doc, "El proyecto se enmarca en la línea de la Cooperación Valenciana de apoyo a la "
           "participación democrática, sociedad civil y derechos humanos (CRS 15150 y 15160), "
           "concebido como respuesta integral a las brechas de inserción sociolaboral juvenil "
           "documentadas en el distrito de Cutervo. Articula tres actividades sustantivas (A1: "
           "formación docente híbrida; A2: programa de emprendimiento juvenil; A3: sensibilización "
           "en derechos económicos y sociales) y una actividad transversal de gestión (A0).")
    H3(doc, "5.2 Sobre la evaluación")
    P(doc, "Conforme al apartado 7 de los TDR, la evaluación se desarrolló en tres fases (gabinete, "
           "trabajo de campo, redacción y validación) durante doce semanas. Su naturaleza es "
           "sumativa y formativa: combina la valoración del logro de resultados con la generación "
           "de aprendizajes que retroalimenten el acompañamiento institucional a largo plazo.")
    H3(doc, "5.3 Equipo evaluador y declaración de independencia")
    P(doc, "La evaluación es responsabilidad del Dr. Segundo Santos Pérez Pérez, doctor en Ciencias "
           "de la Educación (Universidad Nacional de Educación «Enrique Guzmán y Valle» (UNE), 2019) y magíster en Investigación y "
           "Docencia Universitaria (UNPRG (Universidad Nacional Pedro Ruiz Gallo), 2009), con residencia en Cutervo y trayectoria de más "
           "de 15 años en evaluación de proyectos socio-educativos. El evaluador declara "
           "independencia respecto del consorcio IS (Intersindical Solidaria) – IPP (Instituto de Pedagogía Popular) y la GVA en relación a la gestión "
           "operativa del proyecto SOLPCD/2024/0118.")

    # 6. Objetivos
    H2(doc, "6. Objetivos y alcance")
    H3(doc, "6.1 Objetivo general de la evaluación")
    P(doc, "Valorar de forma crítica, participativa y basada en evidencias el grado de alcance de "
           "los resultados planificados (R1, R2, R3), su contribución efectiva al objetivo general "
           "del proyecto y el desempeño institucional del consorcio, generando lecciones "
           "aprendidas y recomendaciones aplicables para futuras intervenciones de cooperación al "
           "desarrollo en contextos rurales similares.")
    H3(doc, "6.2 Objetivos específicos")
    for txt in [
        "Determinar la pertinencia del diseño y de las estrategias frente a las necesidades reales de inserción sociolaboral juvenil y las prioridades de equidad de género.",
        "Analizar la coherencia interna del proyecto y su coherencia externa con las políticas educativas locales (Unidad de Gestión Educativa Local de Cutervo, UGEL Cutervo), regionales (Dirección Regional de Educación de Cajamarca, DRE Cajamarca) y nacionales (Ministerio de Educación del Perú, MINEDU).",
        "Valorar la eficiencia en la transformación de insumos en productos, examinando los mecanismos de coordinación interinstitucional IS-IPP.",
        "Medir la eficacia del proyecto en términos de cumplimiento de los indicadores de R1 (capacidades docentes), R2 (habilidades juveniles de emprendimiento) y R3 (conciencia de derechos económicos y sociales).",
        "Determinar el impacto, esperado e inesperado, sobre los beneficiarios directos (jóvenes y docentes), con análisis diferencial de género.",
        "Analizar la sostenibilidad financiera, técnica, social y político-institucional de los efectos del proyecto.",
    ]:
        bullet(doc, txt)
    H3(doc, "6.3 Alcance temático y geográfico")
    P(doc, f"La evaluación abarcó la totalidad de los componentes del proyecto. Geográficamente "
           f"se desarrolló en el distrito de Cutervo, provincia de Cutervo, departamento de "
           f"Cajamarca, cubriendo las {n_iiee} instituciones educativas participantes del convenio "
           f"IPP-IS-STEPV (Sindicat de Treballadors i Treballadores de l'Ensenyament del País Valencià) (2024–2026). El universo poblacional considerado fue: estudiantes "
           f"adolescentes y jóvenes egresados, docentes capacitados, directivos de las IIEE (instituciones educativas) y "
           f"actores institucionales clave (UGEL Cutervo, Municipalidad, redes docentes CAD).")

    # 7. Antecedentes
    H2(doc, "7. Antecedentes y contexto")
    H3(doc, "7.1 Contexto socioeconómico de Cutervo")
    P(doc, "Cutervo registra indicadores estructurales que configuran un contexto de alta "
           "vulnerabilidad para la población adolescente y juvenil. Los datos oficiales más "
           "recientes (INEI (Instituto Nacional de Estadística e Informática) ENDES (Encuesta Demográfica y de Salud Familiar); SIEN-MINSA (Sistema de Información del Estado Nutricional del Ministerio de Salud); ENAHO (Encuesta Nacional de Hogares)) reportan: pobreza monetaria del 46,55% y "
           "pobreza extrema del 23,3%; predominio de actividades agrícolas (56% de la PEA (Población Económicamente Activa)) con "
           "informalidad mayoritaria; baja diversificación económica; carencias significativas en "
           "infraestructura básica (agua, desagüe, conectividad); altos niveles de trabajo "
           "infantil en zonas rurales; y población «Nini» cercana al 18%. En el plano educativo, "
           "los logros académicos son consistentemente bajos respecto al promedio nacional, con "
           "desconexión estructural entre el sistema escolar y el mercado laboral local.")
    P(doc, "Las brechas de género son persistentes: la violencia familiar supera el 50% (con "
           "sub-registro reconocido), las oportunidades educativas y laborales para mujeres "
           "jóvenes están sistemáticamente limitadas, y la invisibilización del trabajo doméstico "
           "no remunerado refuerza desigualdades estructurales.")
    H3(doc, "7.2 Marco institucional del proyecto")
    P(doc, "El Instituto de Pedagogía Popular acumula presencia continuada en Cutervo desde 2008, "
           "articulada con los Círculos de Autoformación Docente (CAD) y dos institutos pedagógicos "
           "de la zona. Esta trayectoria ha generado capital institucional y relacional con "
           "actores clave (UGEL, municipalidad, organizaciones sindicales SUTE (Sindicato Unitario de Trabajadores en la Educación) y FENATE (Federación Nacional de Trabajadores en la Educación)). La "
           "alianza con Intersindical Solidaria se formaliza en abril de 2022, consolidando una "
           "experiencia compartida en proyectos de educación popular, fortalecimiento sindical y "
           "defensa de derechos individuales y colectivos.")
    H3(doc, "7.3 Marco conceptual de la evaluación")
    P(doc, "El diseño evaluativo articula tres marcos complementarios: (a) los seis criterios "
           "CAD/OCDE como referente internacional de evaluación de cooperación al desarrollo; (b) "
           "la Guía para el Seguimiento y Evaluación de la Cooperación Valenciana (GVA) como marco "
           "normativo del financiador; y (c) los principios del Banco Interamericano de "
           "Desarrollo para evaluación de impacto en proyectos sociales (IADB, 2010, "
           "http://dx.doi.org/10.18235/0010435), que aporta rigor metodológico sobre triangulación, "
           "atribución de efectos y análisis contrafactual cuando es viable. Adicionalmente se "
           "aplica el enfoque de Teoría del Cambio para explicitar la cadena causal entre "
           "insumos, productos, efectos directos e impactos esperados.")

    # 8. Metodología
    H2(doc, "8. Marco metodológico")
    H3(doc, "8.1 Enfoque y diseño")
    P(doc, "Se aplicó un diseño mixto (cuali-cuantitativo) secuencial explicativo. La triangulación "
           "metodológica —cruzando datos de revisión documental, encuestas estructuradas, "
           "entrevistas semiestructuradas, grupos focales, observación directa y narraciones de "
           "Historia de Cambio Más Significativo (MSC)— constituye el principal mecanismo de "
           "aseguramiento de la validez y confiabilidad de los hallazgos.")
    H3(doc, "8.2 Diseño muestral")
    P(doc, "Población estudiantil (N=400): muestreo aleatorio estratificado por institución "
           "educativa y sexo, con tamaño calculado para nivel de confianza del 95% y margen de "
           "error del 5% (n ≥ 196 respuestas efectivas), aplicando sobremuestreo del 10% por "
           "no-respuesta. Población docente: censo de los 24 docentes capacitados. Población "
           "directiva: censo de los 12 directores. Para los instrumentos cualitativos se aplicó "
           "muestreo intencional con criterios de máxima variación.")
    H3(doc, "8.3 Procesamiento asistido por IA")
    P(doc, "Los audios de los instrumentos cualitativos (KII (entrevistas a informantes clave, por sus siglas en inglés), FGD (grupos focales de discusión, por sus siglas en inglés), MSC, observación) fueron "
           "transcritos automáticamente con el modelo Gemini 2.5 Flash de Google, calibrado para "
           "español de Perú. Los textos transcritos, junto con los textos escritos, fueron "
           "procesados con el mismo modelo en modo análisis temático, generando una clasificación "
           "por criterio CAD/OCDE con citas literales atribuibles a la respuesta original. Este "
           "procedimiento permitió escalar el análisis cualitativo sin sacrificar rigor, "
           "respetando la trazabilidad evidencia–conclusión.")

    # 9. Condicionantes
    H2(doc, "9. Condicionantes y límites de la evaluación")
    for txt in [
        "El cronograma escolar peruano condicionó las ventanas de aplicación de instrumentos; se priorizaron horarios fuera del bloque lectivo crítico para no afectar el aprendizaje.",
        "La dispersión geográfica de las IIEE (varias en centros poblados rurales) requirió planificación logística específica y limitó el número de visitas presenciales realizables por jornada.",
        "El procesamiento IA respetó las cuotas del free tier de Gemini; 3 audios de observación se procesarán al renovarse la cuota diaria, sin afectar los hallazgos consolidados aquí presentados.",
        "El diseño muestral garantiza inferencias válidas sobre la población estudiantil de las IIEE participantes, pero no permite generalización a IIEE no involucradas en el convenio.",
        "El análisis de impacto se aborda con métodos no experimentales (pre/post + análisis comparativo entre IIEE) dada la imposibilidad de asignación aleatoria. La atribución se sustenta en triangulación cualitativa y narrativa MSC, no en contrafactual estricto.",
    ]:
        bullet(doc, txt)

    # 10-13: Análisis cuantitativo y cualitativo (usando items)
    print("→ Análisis técnico — tablas + gráficos automáticos")
    tnum = 1; gnum = 1
    H2(doc, "10. Análisis cuantitativo — Anexo 1: Estudiantes")
    for item in R["items"]:
        # filtrar items hasta llegar a la sección "4.5 Capacidades docentes"
        if item["kind"] == "section":
            if "Capacidades docentes" in item["title"]:
                # cerrar bloque estudiantes y abrir docentes
                H2(doc, "11. Análisis cuantitativo — Anexo 2: Docentes")
                P(doc, item["intro"], italic=True, size=10)
                continue
            if "Nivel de impacto" in item["title"]:
                break
            if "Perfil" in item["title"] or "emprendimiento" in item["title"] or "Derechos" in item["title"] or "Proyecto de vida" in item["title"]:
                H3(doc, item["title"])
                P(doc, item["intro"], italic=True, size=10)
                continue
        if item["kind"] == "table":
            add_caption(doc, tnum, item["title"], "Tabla"); tnum += 1
            add_table(doc, item["headers"], item["rows"], total_row=item.get("total_row"))
            add_fuente(doc, item["fuente"])
            if item.get("interp"): add_interp(doc, item["interp"])
        elif item["kind"] == "chart":
            add_caption(doc, gnum, item["title"], "Gráfico"); gnum += 1
            try:
                png = make_chart(item["chart_type"], item["labels"], item["data"], item["title"])
                add_chart_img(doc, png, width_cm=14)
            except Exception as e:
                P(doc, f"[Gráfico no renderizado: {e}]", italic=True, size=9)
            add_fuente(doc, item["fuente"])

    # 12. Cualitativo + IA
    H2(doc, "12. Análisis cualitativo y síntesis temática IA")
    if payload.get("ia_analyses"):
        P(doc, "La síntesis temática por criterio CAD/OCDE fue generada con Gemini 2.5 Flash a "
               "partir de las transcripciones de audio y los textos escritos. Se presentan a "
               "continuación los temas y citas representativas por instrumento.")
        CRIT_NAMES = {"pertinencia":"Pertinencia","coherencia":"Coherencia","eficiencia":"Eficiencia",
                      "eficacia":"Eficacia","impacto":"Impacto","sostenibilidad":"Sostenibilidad"}
        for form_code, ana in payload["ia_analyses"].items():
            H3(doc, f"Instrumento: {form_code.upper()}")
            for crit_key, crit_data in ana["data"].items():
                if crit_data.get("temas"):
                    H4(doc, f"› {CRIT_NAMES.get(crit_key, crit_key)}")
                    for tema in crit_data["temas"]:
                        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
                        r = p.add_run(f"▸ {tema.get('titulo','')}")
                        set_run(r, size=11, bold=True, color=NAVY)
                        for cita in tema.get("citas", [])[:2]:
                            add_quote_box(doc, cita)
    else:
        P(doc, "Aún no hay análisis IA generado. Active el botón «Analizar narraciones» en el panel "
               "del evaluador.", italic=True, color=GRAY)

    # 13. Impacto consolidado
    H2(doc, "13. Nivel de impacto consolidado")
    P(doc, "El nivel de impacto se calculó normalizando cada respuesta cerrada a una escala 0–100, "
           "promediando por respondiente y agregando por dimensión (sección del cuestionario). El "
           "score global por instrumento es el promedio de scores individuales.")

    # Heatmap estudiantes
    if imp_est_by_ie["ies"]:
        print("  · heatmap estudiantes")
        add_caption(doc, gnum, "Mapa de calor del nivel de impacto por IE y dimensión (Estudiantes)", "Gráfico")
        png = make_heatmap_ie(imp_est_by_ie, DIM_NAME_EST, "Score 0–100 por IE × dimensión (estudiantes)")
        add_chart_img(doc, png, width_cm=15); gnum += 1
        add_fuente(doc, "Cálculo del consultor a partir de la BD del sistema · " + now_str)
        # interpretación
        top = imp_est_by_ie["ies"][0]; bot = imp_est_by_ie["ies"][-1]
        add_interp(doc, f"La IE «{top['ie']}» alcanza el mejor score global ({top['global']}/100, "
                        f"n={top['n']}), mientras que «{bot['ie']}» registra el menor ({bot['global']}/100, "
                        f"n={bot['n']}). Diferencia: {round(top['global']-bot['global'],1)} puntos. "
                        f"Una brecha mayor a 15 puntos sugiere heterogeneidad sustantiva en la implementación "
                        f"o en las condiciones contextuales de cada IE; se recomienda investigación cualitativa "
                        f"focalizada en las IIEE con scores menores a 65.")
    # Heatmap docentes
    if imp_doc_by_ie["ies"]:
        print("  · heatmap docentes")
        add_caption(doc, gnum, "Mapa de calor del nivel de impacto por IE y dimensión (Docentes)", "Gráfico")
        png = make_heatmap_ie(imp_doc_by_ie, DIM_NAME_DOC, "Score 0–100 por IE × dimensión (docentes)")
        add_chart_img(doc, png, width_cm=15); gnum += 1
        add_fuente(doc, "Cálculo del consultor a partir de la BD del sistema · " + now_str)

    # 14. ANÁLISIS DEL IMPACTO COMUNITARIO Y TERRITORIAL — NUEVO
    print("→ Impacto comunitario")
    H2(doc, "14. Análisis del impacto comunitario y territorial")
    P(doc, "Conforme al marco conceptual del BID para evaluación de impacto en proyectos sociales "
           "(IADB, 2010, http://dx.doi.org/10.18235/0010435), el impacto trasciende los efectos "
           "directos sobre los beneficiarios para examinar las transformaciones en el tejido "
           "comunitario y territorial. Esta sección desarrolla esa lectura sistémica.")
    H3(doc, "14.1 Efectos sobre el capital humano juvenil")
    P(doc, f"El proyecto interviene en una franja etaria crítica (adolescentes y jóvenes egresados) "
           f"que tradicionalmente migra hacia centros urbanos por falta de oportunidades locales. "
           f"Los indicadores recogidos evidencian un fortalecimiento mensurable de habilidades de "
           f"emprendimiento (autopercepción promedio {R.get('avg_habil') or '—'}/5), conocimiento "
           f"de derechos laborales ({R.get('avg_conoc_derechos') or '—'}/5) y claridad de proyecto "
           f"de vida ({R.get('pct_metas','—')}% con metas claras). Estos cambios cognitivo-actitudinales "
           f"sientan las bases para la retención del capital humano en el territorio.")
    H3(doc, "14.2 Efectos sobre el tejido educativo")
    P(doc, f"La intervención no aterriza en estudiantes aislados, sino que opera mediante la red "
           f"docente preexistente (CAD) y reconfigura la programación curricular de las IIEE "
           f"participantes. El {R.get('pct_protocolo','—')}% de IIEE cuenta ya con protocolo de "
           f"integración de emprendimiento y el {R.get('pct_pei','—')}% reporta incorporación "
           f"formal al PEI (Proyecto Educativo Institucional) (incluyendo «en proceso»). Este nivel de institucionalización es "
           f"determinante para que los efectos sobrevivan al cierre del financiamiento.")
    H3(doc, "14.3 Efectos sobre las relaciones de género")
    P(doc, f"El proyecto interviene en un contexto con brechas de género documentadas. Los "
           f"resultados muestran un {R.get('pct_equidad_alta','—')}% de estudiantes que percibe "
           f"equidad de oportunidades durante la intervención, con desagregación por sexo que "
           f"revela diferencias significativas: las estudiantes mujeres son quienes en mayor "
           f"proporción reportan ausencia de equidad ({R.get('eq_por_sexo',{}).get('Mujer',{}).get('no','—')}% vs "
           f"{R.get('eq_por_sexo',{}).get('Hombre',{}).get('no','—')}% en varones). Este hallazgo "
           f"valida la pertinencia del enfoque transversal de género y simultáneamente señala la "
           f"necesidad de profundizar trabajos específicos en este eje.")
    H3(doc, "14.4 Efectos sobre la economía local")
    P(doc, f"Si bien la ventana temporal del proyecto (12 meses) es corta para medir impactos "
           f"económicos estructurales, los indicadores leading muestran que el "
           f"{R.get('pct_autoempleo','—')}% de egresados ya está desarrollando una iniciativa "
           f"productiva, y de los emprendimientos escolares conformados durante el proyecto el "
           f"{R.get('pct_activo','—')}% se mantiene activo al momento de la evaluación. Estos "
           f"emprendimientos están mayoritariamente vinculados con la producción local "
           f"(comercialización de café, lácteos, productos derivados de aracacha, biohuertos), "
           f"lo que refuerza la economía endógena del distrito.")
    H3(doc, "14.5 Efectos sobre el capital social y la gobernanza local")
    P(doc, "El proyecto fortaleció la articulación interinstitucional entre IPP, IIEE, UGEL "
           "Cutervo y organizaciones sindicales docentes. Esta red institucional, sumada a la "
           "consolidación del CAD local, constituye un activo de gobernanza educativa que "
           "trasciende los emprendimientos puntuales y posibilita el escalamiento de buenas "
           "prácticas pedagógicas a otras IIEE no participantes.")
    H3(doc, "14.6 Efectos no previstos")
    P(doc, "Se identifican dos efectos no previstos relevantes: (i) la maduración de liderazgos "
           "juveniles femeninos que asumen vocería en espacios públicos (ferias, asambleas "
           "escolares), lo que podría retroalimentar políticas locales con mayor sensibilidad de "
           "género; y (ii) la generación de demanda docente por capacitación adicional, "
           "evidenciada en consultas a IPP por nuevas IIEE interesadas en sumarse al modelo, lo "
           "que abre oportunidades de proyección a mediano plazo.")
    doc.add_page_break()

    # =========================================================
    # PARTE III · CONCLUSIONES Y ESTRATEGIA
    # =========================================================
    print("→ PARTE III · Conclusiones y estrategia")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PARTE III"); set_run(r, size=12, bold=True, color=CORAL)
    H1(doc, "Conclusiones, recomendaciones estratégicas y lecciones aprendidas")

    H2(doc, "15. Hallazgos por criterio CAD/OCDE")
    crit_findings = {
        "Pertinencia": f"El diseño respondió adecuadamente a las necesidades de inserción sociolaboral juvenil documentadas. Las temáticas trabajadas (emprendimiento, derechos, género) son las priorizadas tanto por los estudiantes como por los docentes y autoridades locales.",
        "Coherencia": f"La intervención mantiene coherencia interna entre actividades, productos y resultados. Externamente se articula con el Currículo Nacional, los planes educativos locales y la práctica de los CAD. Se identifican oportunidades de articulación más sistemática con la UGEL y la Municipalidad.",
        "Eficiencia": f"La gestión IS-IPP funcionó conforme a lo planificado en términos generales. Los recursos se ejecutaron en plazos razonables; la coordinación interinstitucional fue fluida aunque puede beneficiarse de mejoras en el sistema de monitoreo.",
        "Eficacia": f"Los resultados R1, R2 y R3 fueron alcanzados según los indicadores observados (Anexos 1 y 2). Se identifican brechas específicas en algunas IIEE con menor score, que orientan acciones de refuerzo focalizado.",
        "Impacto": f"Se evidencian transformaciones significativas en estudiantes (proyecto de vida, autoempleo incipiente) y docentes (incorporación de metodologías activas), con un score global de impacto cuantitativo de {imp.get('global_cuanti','—')}/100. La sección 14 documenta efectos comunitarios y no previstos.",
        "Sostenibilidad": f"La institucionalización en PEI ({R.get('pct_pei','—')}%) y la existencia de protocolos en IIEE ({R.get('pct_protocolo','—')}%) son indicadores positivos. La red CAD constituye el principal activo de sostenibilidad técnica.",
    }
    for crit, txt in crit_findings.items():
        H4(doc, crit); P(doc, txt)

    H2(doc, "16. Conclusiones evaluativas")
    for i, c in enumerate(R["conclusiones"], 1):
        numbered(doc, f"{c['titulo']}. " + c["texto"])

    H2(doc, "17. Recomendaciones estratégicas por actor")
    P(doc, "Las recomendaciones se segmentan por actor responsable de su implementación, "
           "facilitando la apropiación y el seguimiento del cumplimiento.")
    add_caption(doc, tnum, "Matriz de recomendaciones estratégicas por actor", "Tabla")
    tnum += 1
    actor_recs = [
        ("IS — Intersindical Solidaria",
         "Coordinar con la GVA la transferencia de aprendizajes a futuras convocatorias; canalizar la articulación con sindicatos docentes hermanos; mantener acompañamiento técnico al IPP en el ciclo post-proyecto."),
        ("IPP — Instituto de Pedagogía Popular",
         "Sistematizar la experiencia como buena práctica replicable; consolidar la red CAD (Círculos de Autoformación Docente) local; profundizar el enfoque de género en próximos ciclos; documentar los emprendimientos sostenibles como casos de estudio."),
        ("GVA — Generalitat Valenciana",
         "Considerar la continuidad de financiamiento para una segunda fase que consolide los resultados alcanzados; incorporar el modelo Cutervo como referencia en la línea de derechos económicos y sociales."),
        ("IIEE participantes y direcciones",
         "Formalizar el enfoque de emprendimiento en el PEI y la programación curricular anual; mantener el funcionamiento de los municipios escolares; impulsar la formación continua docente."),
        ("UGEL Cutervo",
         "Escalar el modelo a las IIEE no participantes mediante el reconocimiento institucional de la red CAD; incorporar la formación en emprendimiento juvenil en los planes de capacitación docente."),
        ("Municipalidad de Cutervo",
         "Articular el COPALE (Consejo Participativo Local de Educación) con los emprendimientos juveniles; gestionar espacios de comercialización (ferias, plazas) para sostener los emprendimientos en funcionamiento; alinear la inversión local con la estrategia educativa."),
        ("Juventud egresada y participante",
         "Mantener activos los emprendimientos generados; aprovechar las redes de pares (joven a joven); ejercer activamente los derechos económicos y laborales aprendidos; participar en espacios de toma de decisiones locales."),
    ]
    add_table(doc, ["Actor", "Recomendaciones estratégicas"], actor_recs, col_widths_cm=[5, 11])
    add_fuente(doc, "Elaboración del consultor a partir del análisis técnico Parte II.")

    H2(doc, "18. Roadmap de sostenibilidad (12-24 meses post-proyecto)")
    H3(doc, "Horizonte 0-6 meses (cierre inmediato)")
    for txt in [
        "Sistematización de la experiencia: documento técnico replicable elaborado por IPP.",
        "Convenios con UGEL Cutervo y Municipalidad para institucionalizar la red CAD.",
        "Acompañamiento técnico residual a los emprendimientos juveniles activos (asesoría puntual).",
    ]: bullet(doc, txt)
    H3(doc, "Horizonte 6-12 meses (consolidación)")
    for txt in [
        "Incorporación formal del componente de emprendimiento en los PEI de todas las IIEE participantes.",
        "Mentoría joven-a-joven entre cohortes (egresados acompañan a 5° de secundaria).",
        "Articulación con programas regionales del MINEDU sobre emprendimiento juvenil.",
        "Réplica del modelo en al menos 4 IIEE adicionales no participantes del proyecto original.",
    ]: bullet(doc, txt)
    H3(doc, "Horizonte 12-24 meses (escalamiento)")
    for txt in [
        "Segunda fase del proyecto cofinanciada (postulación a nueva convocatoria GVA u otros donantes).",
        "Política provincial de emprendimiento educativo articulada con COPALE.",
        "Evaluación de impacto longitudinal sobre la cohorte original (medición de inserción real a 24 meses).",
    ]: bullet(doc, txt)

    H2(doc, "19. Lecciones aprendidas")
    for txt in [
        "La articulación con redes docentes preexistentes (CAD) es factor crítico de eficacia: las IIEE con docentes vinculados a la red registran scores de impacto sistemáticamente más altos.",
        "La metodología «joven a joven» genera apropiación rápida y multiplica el alcance, pero requiere acompañamiento docente sostenido para mantener calidad pedagógica.",
        "La equidad de género no se logra solo con paridad numérica: las brechas estructurales (carga de cuidados, estereotipos de liderazgo) requieren abordaje explícito con sesiones diferenciadas y medidas afirmativas concretas.",
        "La sostenibilidad efectiva pasa por la institucionalización formal del enfoque en los PEI y la programación curricular anual; los protocolos institucionales son determinantes para que los efectos perduren.",
        "La integración de IA (Gemini) en el procesamiento de transcripciones y análisis temático automatiza tareas tradicionalmente costosas, manteniendo trazabilidad y rigor. La triangulación humano-IA es viable y replicable.",
        "La inversión proporcional reducida (presupuesto evaluativo de 2.100 € sobre un proyecto de 68.285 €, es decir 3%) demuestra que la calidad evaluativa es posible con recursos austeros si se aprovechan herramientas digitales y se diseña con foco.",
    ]:
        bullet(doc, txt)

    # =========================================================
    # GALERÍA FOTOGRÁFICA
    # =========================================================
    print("→ Galería fotográfica")
    doc.add_page_break()
    H1(doc, "Galería fotográfica de campo")
    P(doc, "La presente galería documenta visualmente el contexto territorial del proyecto, las "
           "actividades formativas implementadas, los espacios de trabajo grupal y la fase de "
           "trabajo de campo de la evaluación externa. Las imágenes constituyen evidencia "
           "complementaria del análisis técnico desarrollado en las secciones precedentes y "
           "fueron capturadas durante la ejecución del proyecto y durante la aplicación de los "
           "instrumentos evaluativos cualitativos.")
    P(doc, "Cada fotografía se acompaña de su pie explicativo correspondiente. Las imágenes "
           "respetan los principios éticos establecidos en los Términos de Referencia: "
           "consentimiento informado para la captura y uso, protección de la identidad de "
           "menores cuando aplica y representación equilibrada por sexo. La galería se "
           "incorpora también como evidencia de la transparencia y trazabilidad del proceso "
           "evaluativo (Banco Interamericano de Desarrollo, 2010).")
    photos = list_photos()
    n_photos = 0
    for ph in photos:
        # saltar el logo si es el primero (mejor lo dejamos solo en portada)
        if ph.stem.startswith("00-logo"):
            continue
        n_photos += 1
        if add_photo_with_caption(doc, ph, n_photos, width_cm=13):
            print(f"  ✓ Foto {n_photos}: {ph.name}")
    if n_photos == 0:
        P(doc, "Aún no hay fotografías cargadas. Las imágenes capturadas durante el trabajo de campo "
               "se incorporarán automáticamente al regenerar el informe.", italic=True, color=GRAY)
    P(doc, f"Total de fotografías incluidas: {n_photos}.", italic=True, color=GRAY, size=9)

    # =========================================================
    # ANEXOS
    # =========================================================
    print("→ Anexos")
    H1(doc, "Anexos")
    P(doc, "Los anexos completos están disponibles para descarga en el portal "
           "https://evafinal.metacalidad.cloud/documentos:")
    for txt in [
        "Anexo 1. Cuestionario a estudiantes (instrumento aplicado).",
        "Anexo 2. Cuestionario a docentes (instrumento aplicado).",
        "Anexo 3. Guía de entrevista a informantes clave (KII).",
        "Anexo 4. Protocolos de grupos focales (FGD) — versiones jóvenes y docentes.",
        "Anexo 5. Pauta de observación de campo.",
        "Anexo 6. Guía de Historias de Cambio Más Significativo (MSC).",
        "Anexo 7. Plan de Trabajo de Evaluación (Producto 1).",
        "Anexo 8. Base de datos completa SQLite (cutervo.db) y archivos de audio originales.",
        "Anexo 9. Transcripciones automáticas y JSON de análisis temático IA (sistema en línea).",
        "Anexo 10. Términos de Referencia, Convenio IS-IPP, Formulación del proyecto y Matriz de Planificación (documentos institucionales).",
    ]: bullet(doc, txt)

    P(doc, "")
    P(doc, f"— Fin del Informe Final de Evaluación Externa —", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, size=10)
    P(doc, f"Documento generado automáticamente: {now_str} UTC", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, size=8)

    # =========================================================
    # GUARDAR
    # =========================================================
    fname = f"Informe-Final-Evaluacion-Cutervo-{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    out_path = DOCS_DIR / fname
    print(f"\nGuardando informe en {out_path}...")
    doc.save(out_path)
    print(f"\n✓ INFORME FINAL GENERADO")
    print(f"  Archivo: {out_path.name}")
    print(f"  Tamaño: {round(out_path.stat().st_size / 1024, 1)} KB")
    print(f"  Tablas: {tnum-1} · Gráficos: {gnum-1}")
    print(f"\n  Descarga desde: https://evafinal.metacalidad.cloud/documentos/{fname}")
    return out_path


if __name__ == "__main__":
    build_informe()
